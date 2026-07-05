import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(sys.argv[1])
OUTPUT = Path(sys.argv[2])


def clean(value):
    return " ".join((value or "").split())


def split_identifier(value):
    match = re.match(r"^([^ ]+)\s+-\s+(.+)$", value)
    if not match:
        return value, value
    return match.group(1), match.group(2)


document = Document(SOURCE)
blocks = []
for child in document.element.body.iterchildren():
    if isinstance(child, CT_P):
        paragraph = Paragraph(child, document)
        text = clean(paragraph.text)
        if text:
            blocks.append({"kind": "paragraph", "style": paragraph.style.name, "text": text})
    elif isinstance(child, CT_Tbl):
        table = Table(child, document)
        rows = [[clean(cell.text) for cell in row.cells] for row in table.rows]
        if rows:
            blocks.append({"kind": "table", "rows": rows})


features = []
stories = []
current_feature = None
current_story = None
current_section = None

for block in blocks:
    if block["kind"] == "paragraph":
        style = block["style"]
        text = block["text"]
        if style == "Heading 3" and text.startswith("FEAT-"):
            feature_id, title = split_identifier(text)
            current_feature = {
                "id": feature_id,
                "title": title,
                "description": "",
            }
            features.append(current_feature)
            current_story = None
            current_section = None
        elif style == "Heading 4" and text.startswith("US-"):
            story_id, title = split_identifier(text)
            current_story = {
                "id": story_id,
                "featureId": current_feature["id"] if current_feature else None,
                "title": title,
                "persona": "",
                "need": "",
                "benefit": "",
                "description": "",
                "sourceTrace": "",
                "priority": "High",
                "complexity": "Medium",
                "dependencies": [],
                "acceptanceCriteria": [],
                "qaFocus": [],
                "tasks": [],
            }
            stories.append(current_story)
            current_section = None
        elif current_story and style == "Intense Quote":
            current_section = text
        elif current_story and style.startswith("List"):
            if current_section == "Acceptance Criteria":
                current_story["acceptanceCriteria"].append(text)
            elif current_section == "QA Focus":
                current_story["qaFocus"].append(text)
    elif current_story and block["kind"] == "table":
        rows = block["rows"]
        if rows and rows[0][:2] == ["Field", "Details"]:
            fields = {row[0]: row[1] for row in rows[1:] if len(row) >= 2}
            current_story["persona"] = fields.get("User Story", "")
            current_story["need"] = fields.get("I want", "")
            current_story["benefit"] = fields.get("So that", "")
            current_story["description"] = " ".join(
                value for value in [current_story["persona"], current_story["need"], current_story["benefit"]] if value
            )
            current_story["sourceTrace"] = fields.get("Source Trace", "")
            priority_complexity = fields.get("Priority / Complexity", "High / Medium").split("/")
            current_story["priority"] = clean(priority_complexity[0])
            current_story["complexity"] = clean(priority_complexity[1]) if len(priority_complexity) > 1 else "Medium"
            current_story["dependencies"] = [clean(item) for item in fields.get("Dependencies", "").split(",") if clean(item)]
        elif rows and rows[0][:3] == ["Task ID", "Owner", "Task / Expected Output"]:
            current_story["tasks"] = [
                {"id": row[0], "ownerRole": row[1], "title": row[2]}
                for row in rows[1:]
                if len(row) >= 3 and row[0]
            ]


feature_story_counts = {feature["id"]: 0 for feature in features}
for story in stories:
    if story["featureId"] in feature_story_counts:
        feature_story_counts[story["featureId"]] += 1
for feature in features:
    feature["storyCount"] = feature_story_counts[feature["id"]]


document_map = []
heading_stack = []
for block in blocks:
    if block["kind"] == "paragraph" and block["style"].startswith("Heading"):
        level_match = re.search(r"(\d+)$", block["style"])
        level = int(level_match.group(1)) if level_match else 1
        heading_stack = heading_stack[: level - 1]
        heading_stack.append(block["text"])
        document_map.append({
            "kind": "heading",
            "level": level,
            "text": block["text"],
            "path": list(heading_stack),
        })
    elif block["kind"] == "paragraph":
        document_map.append({
            "kind": "text",
            "style": block["style"],
            "text": block["text"],
            "path": list(heading_stack),
        })
    else:
        document_map.append({
            "kind": "table",
            "rows": block["rows"],
            "path": list(heading_stack),
        })


payload = {
    "source": {
        "title": "Project Atlas Payroll User Stories",
        "filename": SOURCE.name,
        "scope": ["Basic Pay", "Company Loan", "Bonus Setup", "De Minimis Ceiling Setup"],
        "paragraphCount": len(document.paragraphs),
        "tableCount": len(document.tables),
    },
    "epic": {
        "id": "EPIC-PAYROLL",
        "title": "Atlas Payroll Phase 2",
        "description": "Deliver auditable payroll configuration and computation for Basic Pay, Company Loans, Bonus Setup, De Minimis ceilings, and their cross-module integration.",
    },
    "features": features,
    "stories": stories,
    "document": document_map,
}

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
OUTPUT.write_text(json.dumps(payload, indent=2, ensure_ascii=True), encoding="utf-8")
print(json.dumps({
    "output": str(OUTPUT),
    "features": len(features),
    "stories": len(stories),
    "tasks": sum(len(story["tasks"]) for story in stories),
    "acceptanceCriteria": sum(len(story["acceptanceCriteria"]) for story in stories),
    "qaFocus": sum(len(story["qaFocus"]) for story in stories),
    "documentBlocks": len(document_map),
}, indent=2))
